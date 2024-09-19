from docx import Document

def generate_brd_docx(brd_content):
    doc = Document()
    doc.add_heading('Business Requirements Document', 0)
    doc.add_paragraph(brd_content)
    doc.save("BRD_Document.docx")